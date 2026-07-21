from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument


class DocxExtractor:
    """Extracts text from Word (.docx) documents.

    .docx has no fixed pagination outside of a rendering engine, so the
    whole document is treated as a single page -- consistent with how
    TextExtractor handles .txt/.md and with what the downstream chunker
    expects (a list of page dicts).
    """

    def extract(self, path: Path) -> list[dict]:
        document = DocxDocument(str(path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs]
        parts = [part for part in parts if part]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n\n".join(parts).strip()
        return [{"file": path.name, "page": 1, "text": text}]


docx_extractor = DocxExtractor()
