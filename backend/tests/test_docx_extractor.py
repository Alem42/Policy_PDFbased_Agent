from __future__ import annotations

import shutil
import tempfile
from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document as DocxDocument

from app.modules.documents.docx_extractor import DocxExtractor


@pytest.fixture
def tmp_path():
    """See test_pdf_extractor.py: the shared pytest-of-<user> temp root is
    permission-denied in this environment, so use our own throwaway dir."""
    path = Path(tempfile.gettempdir()) / f"docx_extractor_test_{uuid4().hex}"
    path.mkdir(parents=True, exist_ok=True)
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


def _build_docx(
    tmp_path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None
) -> Path:
    document = DocxDocument()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values, strict=True):
                cell.text = value
    path = tmp_path / "test.docx"
    document.save(path)
    return path


def test_extracts_paragraph_text_as_single_page(tmp_path):
    docx_path = _build_docx(tmp_path, ["First paragraph.", "Second paragraph."])

    pages = DocxExtractor().extract(docx_path)

    assert len(pages) == 1
    assert pages[0]["page"] == 1
    assert "First paragraph." in pages[0]["text"]
    assert "Second paragraph." in pages[0]["text"]


def test_skips_empty_paragraphs(tmp_path):
    docx_path = _build_docx(tmp_path, ["Real content.", "", "   "])

    pages = DocxExtractor().extract(docx_path)

    assert pages[0]["text"] == "Real content."


def test_includes_table_text(tmp_path):
    docx_path = _build_docx(
        tmp_path,
        ["Intro paragraph."],
        table_rows=[["Name", "Value"], ["Budget", "$5m"]],
    )

    pages = DocxExtractor().extract(docx_path)

    assert "Name | Value" in pages[0]["text"]
    assert "Budget | $5m" in pages[0]["text"]
